#!/usr/bin/env python3
"""
Report Generator - Advanced Edition
Converts Markdown text into a professionally formatted Word (.docx) document.
Updated: Fixed 'IndexError' by applying color AFTER text is added.
"""

import io
import re
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.helpers import safe_log

def generate_word_report(markdown_content: str, title: str, casino_mode: bool) -> bytes:
    """Converts markdown string to docx bytes."""
    try:
        doc = Document()
        
        # Metadata
        doc.core_properties.title = title
        doc.core_properties.subject = "Casino Audit" if casino_mode else "YMYL Audit"
        
        # Default Style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Header
        doc.add_heading(title, 0)
        
        # Parse content line by line
        for line in markdown_content.split('\n'):
            line = line.strip()
            if not line: continue
            
            if line.startswith('### '):
                doc.add_heading(line[4:], 2)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 1)
            elif line.startswith('# '):
                doc.add_heading(line[2:], 0)
            elif line.startswith('---'):
                p = doc.add_paragraph('_' * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(200, 200, 200)
                else:
                    p.add_run().font.color.rgb = RGBColor(200, 200, 200)
            elif line.startswith('> '):
                # Blockquote (Used for Translations)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                _add_formatted_text(p, line[2:], is_translation=True)
            elif line.startswith('- '):
                # Bullet points
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_text(p, line[2:])
            else:
                # Standard paragraph
                p = doc.add_paragraph()
                _add_formatted_text(p, line)
                
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        safe_log(f"Reporter Error: {e}", "ERROR")
        # Return a basic error document instead of crashing
        return b""

def _add_formatted_text(paragraph, text, is_translation=False):
    """
    Parses markdown-style formatting within a paragraph.
    Supports **bold** and _italics_.
    """
    # 1. Determine Color (But don't apply it yet)
    target_color = None
    if '🔴' in text or 'Critical' in text:
        target_color = RGBColor(231, 76, 60) # Red
    elif '🟠' in text or 'High' in text:
        target_color = RGBColor(230, 126, 34) # Orange
        
    # 2. Split by bold markers first: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # This is bold
            clean_text = part[2:-2]
            run = paragraph.add_run(clean_text)
            run.bold = True
            
            if is_translation: run.italic = True
            if target_color: run.font.color.rgb = target_color
        else:
            # Check for italics inside non-bold parts: _text_
            italic_parts = re.split(r'(_.*?_)', part)
            for sub_part in italic_parts:
                run = None
                if sub_part.startswith('_') and sub_part.endswith('_'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    if sub_part: # Only add if not empty
                        run = paragraph.add_run(sub_part)
                        if is_translation: run.italic = True
                
                # Apply color to this run if needed
                if run and target_color:
                    run.font.color.rgb = target_color
